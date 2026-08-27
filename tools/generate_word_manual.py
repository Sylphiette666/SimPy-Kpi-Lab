# ruff: noqa: E501

from __future__ import annotations

from collections.abc import Iterable, Sequence
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "SimPy_KPI_Lab_详细解释说明.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
TEAL = "2A7F83"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
PALE_TEAL = "EAF5F5"
PALE_YELLOW = "FFF7D6"
PALE_RED = "FCE8E6"
MID_GRAY = "6B7280"
LIGHT_GRAY = "F5F7FA"
BORDER = "C7D2E0"
WHITE = "FFFFFF"
BLACK = "1F2937"


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, ascii_font: str = "Calibri",
                 east_asia_font: str = "Microsoft YaHei") -> None:
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_language(paragraph, language: str = "zh-CN") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    lang = ppr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        ppr.append(lang)
    lang.set(qn("w:eastAsia"), language)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for margin, value in (("top", top), ("start", start),
                          ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color: str = BORDER, size: int = 6) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360) -> None:
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.first_child_found_in("w:tblW")
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(width_dxa))
    tblw.set(qn("w:type"), "dxa")
    layout = tblpr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_repeat_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def prevent_row_split(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trpr.append(cant_split)


def set_cell_width(cell, inches: float) -> None:
    width = Inches(inches)
    cell.width = width
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(int(width / 635)))
    tcw.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 8.5,
                  color: str = BLACK, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    set_paragraph_language(paragraph)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]],
              widths: Sequence[float], *, font_size: float = 8.5,
              header_fill: str = LIGHT_BLUE) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_header(header)
    for index, (cell, label, width) in enumerate(zip(header.cells, headers, widths, strict=True)):
        set_cell_width(cell, width)
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell)
        set_cell_margins(cell)
        set_cell_text(cell, label, bold=True, size=font_size, color=NAVY,
                      align=WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for cell, value, width in zip(row.cells, row_data, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_borders(cell)
            set_cell_margins(cell)
            set_cell_text(cell, str(value), size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(run_color)
    run_props.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run_props.append(fonts)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_props)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])
    set_run_font(run, size=8, color=MID_GRAY)


def add_toc(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开文档时自动更新；如未更新，请在 Word 中按 Ctrl+A，再按 F9。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])
    set_run_font(run, size=10, color=MID_GRAY)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None,
             color: str = BLACK, keep_with_next: bool = False) -> object:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.widow_control = True
    set_paragraph_language(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True, color=color)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=color)
    return paragraph


def add_bullets(doc: Document, items: Sequence[str], level: int = 0) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        set_paragraph_language(paragraph)
        run = paragraph.add_run(item)
        set_run_font(run)


def create_numbering_instance(doc: Document, style_name: str = "List Number") -> int:
    style = doc.styles[style_name]
    style_num_pr = style._element.pPr.numPr
    base_num_id = int(style_num_pr.numId.val)
    numbering = doc.part.numbering_part.element
    base_num = next(
        item
        for item in numbering.findall(qn("w:num"))
        if int(item.get(qn("w:numId"))) == base_num_id
    )
    abstract_num_id = int(base_num.find(qn("w:abstractNumId")).get(qn("w:val")))
    existing_ids = [
        int(item.get(qn("w:numId")))
        for item in numbering.findall(qn("w:num"))
    ]
    new_num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_num_id


def add_numbered(doc: Document, items: Sequence[str]) -> None:
    num_id = create_numbering_instance(doc)
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        ppr = paragraph._p.get_or_add_pPr()
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            ppr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        direct_num_id = OxmlElement("w:numId")
        direct_num_id.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(direct_num_id)
        set_paragraph_language(paragraph)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_code_block(
    doc: Document,
    text: str,
    *,
    font_size: float = 8.3,
    trailing_space: bool = True,
) -> object:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_borders(cell, color="D7DEE8", size=5)
    set_cell_margins(cell, top=100, start=140, bottom=100, end=140)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.paragraph_format.keep_together = False
    run = paragraph.add_run(text.rstrip())
    set_run_font(run, size=font_size, color="273444", ascii_font="Consolas",
                 east_asia_font="DengXian")
    if trailing_space:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_note_box(doc: Document, title: str, body: str, *, kind: str = "info") -> object:
    fill, accent = {
        "info": (PALE_BLUE, BLUE),
        "tip": (PALE_TEAL, TEAL),
        "warn": (PALE_YELLOW, "A06A00"),
        "danger": (PALE_RED, "B42318"),
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color=accent, size=8)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    r1 = p1.add_run(title)
    set_run_font(r1, size=9.5, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    r2 = p2.add_run(body)
    set_run_font(r2, size=9, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def heading(doc: Document, text: str, level: int = 1) -> object:
    paragraph = doc.add_heading(text, level=level)
    if getattr(doc, "_simlab_page_break_pending", False):
        paragraph.paragraph_format.page_break_before = True
        doc._simlab_page_break_pending = False
    set_paragraph_language(paragraph)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.widow_control = True
    return paragraph


def caption(doc: Document, text: str) -> object:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, color=MID_GRAY)
    return paragraph


def add_manual_page_break(doc: Document) -> None:
    doc._simlab_page_break_pending = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for list_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    # Ask Word/LibreOffice to refresh field codes, including the TOC.
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = hp.add_run("SIMPY KPI LAB")
    set_run_font(left, size=8, bold=True, color=BLUE)
    right = hp.add_run("\t详细解释说明")
    set_run_font(right, size=8, color=MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    label = fp.add_run("第 ")
    set_run_font(label, size=8, color=MID_GRAY)
    add_page_field(fp)
    tail = fp.add_run(" 页")
    set_run_font(tail, size=8, color=MID_GRAY)

    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ffp.paragraph_format.space_after = Pt(0)
    run = ffp.add_run("技术手册 · 基于项目版本 0.1.0 · 2026-08-26")
    set_run_font(run, size=8, color=MID_GRAY)

    props = doc.core_properties
    props.title = "SimPy KPI Lab 详细解释说明"
    props.subject = "可统计 KPI、可执行多次实验、可接入 OpenAI API 的 SimPy 模型技术手册"
    props.author = "OpenAI Codex"
    props.keywords = "SimPy, KPI, replication, parameter grid, OpenAI Responses API"
    props.category = "技术说明书"
    props.comments = "依据当前项目源码、示例配置和已验证运行结果编制。"


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SIMPY · EXPERIMENTATION · KPI")
    set_run_font(run, size=9, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("SimPy KPI Lab")
    set_run_font(run, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("详细解释说明")
    set_run_font(run, size=22, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("一个可统计 KPI、可执行多次实验、可选接入 OpenAI Responses API 的可复现离散事件仿真框架")
    set_run_font(run, size=13, color="475569")

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    labels = [
        ("可复现", "稳定命名随机流"),
        ("可比较", "参数网格 × 多次 replication"),
        ("可解释", "本地统计 + 可选 AI 分析"),
    ]
    for cell, (top, bottom) in zip(table.rows[0].cells, labels, strict=True):
        set_cell_width(cell, 2.166)
        set_cell_shading(cell, PALE_BLUE)
        set_cell_borders(cell, color="D5E1EF")
        set_cell_margins(cell, top=140, start=100, bottom=140, end=100)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(top)
        set_run_font(r1, size=11, bold=True, color=BLUE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(bottom)
        set_run_font(r2, size=8.5, color="475569")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("适用对象：模型使用者、业务分析人员、仿真开发者")
    set_run_font(r, size=9.5, color=MID_GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目路径：D:\\文件\\SimPy\\simpy-kpi-lab\\simpy-kpi-lab")
    set_run_font(r, size=9, color=MID_GRAY, ascii_font="Consolas", east_asia_font="DengXian")


def build_document() -> Document:
    doc = Document()
    project_root_display = str(PROJECT_ROOT)
    configure_document(doc)
    add_cover(doc)
    add_manual_page_break(doc)

    heading(doc, "目录", 1)
    add_toc(doc)
    add_note_box(
        doc,
        "阅读建议",
        "第一次使用可先看第 1、3、5、10 章；需要修改模型时重点看第 6、7、12 章。所有“当前实现”均以本项目现有源码为准，“建议增强”单独列出，不代表已经具备。",
        kind="tip",
    )
    heading(doc, "文档范围与基线", 2)
    add_table(
        doc,
        ["项目", "本说明采用的基线"],
        [
            ("代码版本", "pyproject.toml 中的 0.1.0"),
            ("运行环境", "Windows PowerShell；Python 3.11+"),
            ("示例配置", "examples/service_center.yaml"),
            ("示例实验", "3 个容量场景 × 20 次 replication = 60 次仿真"),
            ("核对日期", "2026-08-26（Asia/Shanghai）"),
        ],
        [1.55, 4.95],
        font_size=9,
    )
    add_manual_page_break(doc)

    heading(doc, "1. 这套模型解决什么问题", 1)
    add_body(
        doc,
        "SimPy KPI Lab 把离散事件仿真拆成四个明确层次：业务系统模型、重复实验、KPI 统计和可选的 AI 解释。它适合回答“增加一个服务台是否值得”“等待时间尾部是否改善”“吞吐、服务水平和利用率之间有什么权衡”等容量与流程问题。",
    )
    add_note_box(
        doc,
        "一句话结论",
        "数值由本地 SimPy 与统计代码计算；OpenAI API 只读取配置和汇总结果并生成解释，不参与事件推进，也不会改写 KPI。",
        kind="info",
    )

    heading(doc, "1.1 已实现能力", 2)
    add_table(
        doc,
        ["能力", "当前实现", "边界"],
        [
            ("排队模型", "任意数量的串行工位；每个工位为 FIFO simpy.Resource", "所有实体按相同顺序访问全部工位"),
            ("随机输入", "指数、定值、均匀、三角分布", "时间单位由业务统一约定"),
            ("实验设计", "参数笛卡尔积、每场景多次 replication、可多进程", "未内置自动收敛停止或最优解搜索"),
            ("可复现性", "BLAKE2b 派生 64 位 seed；到达与各工位独立子流", "跨 Python/依赖版本不承诺位级一致"),
            ("KPI", "吞吐、周期、等待、服务水平、利用率、队长、WIP、删失诊断", "完成案例指标仍受右删失影响"),
            ("统计", "均值、样本标准差、标准误、正态近似 CI、最小/最大", "当前不是 Student-t、bootstrap 或配对差值 CI"),
            ("AI 分析", "Responses API + Pydantic Structured Outputs", "显式 --analyze 时才联网；需 API Key 与可用额度"),
        ],
        [1.22, 3.05, 2.23],
        font_size=8.1,
    )

    heading(doc, "1.2 数据流", 2)
    flow = doc.add_table(rows=1, cols=5)
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(flow)
    flow_items = [
        ("① YAML", "配置与实验设计"),
        ("② 场景", "参数网格展开"),
        ("③ SimPy", "逐次 replication"),
        ("④ 本地统计", "KPI 与置信区间"),
        ("⑤ 输出", "JSON/CSV + 可选 AI"),
    ]
    fills = ["EAF2F8", "E8F4F3", "EFF6E8", "FFF4E5", "F3ECFA"]
    for cell, (title, detail), fill in zip(flow.rows[0].cells, flow_items, fills, strict=True):
        set_cell_width(cell, 1.3)
        set_cell_shading(cell, fill)
        set_cell_borders(cell, color="CCD6E2")
        set_cell_margins(cell, top=120, start=70, bottom=120, end=70)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(title)
        set_run_font(r1, size=9, bold=True, color=NAVY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(detail)
        set_run_font(r2, size=7.5, color="475569")
    caption(doc, "图 1  从配置到可解释结果的处理链")

    heading(doc, "1.3 示例业务场景", 2)
    add_body(
        doc,
        "示例是两阶段服务中心：顾客到达后先经过 registration，再经过 specialist，最后离开。到达间隔均值为 5 分钟，登记服务均值为 3 分钟，专家服务使用 4/8/15 分钟的三角分布；实验把 specialist 容量从 2、3、4 逐一比较。",
    )
    add_code_block(doc, "到达 → registration（容量 1） → specialist（容量 2 / 3 / 4） → 完成")

    heading(doc, "2. 项目结构与模块职责", 1)
    add_code_block(
        doc,
        """simpy-kpi-lab/
├─ examples/service_center.yaml   # 可直接运行的两阶段服务中心配置
├─ src/simlab/
│  ├─ config.py                   # YAML、Pydantic 校验、分布抽样
│  ├─ simulation.py               # SimPy 实体、到达、工位与单次运行
│  ├─ rng.py                      # BLAKE2b 命名随机流
│  ├─ kpi.py                      # 事件采集、KPI 口径、指标目录
│  ├─ experiment.py               # 场景展开、多次运行、汇总、CSV/JSON
│  ├─ ai.py                       # 可选 OpenAI Responses API 分析器
│  ├─ cli.py                      # validate / run / analyze 命令
│  └─ __main__.py                 # python -m simlab 入口
├─ tests/                         # 配置、仿真、随机数、实验、AI 测试
├─ outputs/                       # 默认实验输出
├─ pyproject.toml                 # 包、依赖、simlab 命令入口
├─ Dockerfile / Makefile          # 容器与常用命令
└─ README.md                      # 快速说明""",
        font_size=8.1,
    )
    add_table(
        doc,
        ["模块", "关键职责", "典型输入 → 输出"],
        [
            ("config.py", "拒绝未知字段；校验时间、工位、分布和参数网格", "YAML → ProjectConfig"),
            ("simulation.py", "构造环境、资源、顾客过程与到达过程", "SimulationConfig + seed → 单次 metrics"),
            ("kpi.py", "记录事件时间与时间面积；输出 KPI 和 metric_catalog", "事件 → replication KPI"),
            ("experiment.py", "展开场景、派发任务、聚合和保存", "ProjectConfig → results.json/CSV"),
            ("ai.py", "压缩结果、调用 Responses API、校验结构化输出", "summary → AI JSON/Markdown"),
            ("cli.py", "解析命令、输出进度、统一常见错误", "PowerShell 命令 → 退出码与文件"),
        ],
        [1.25, 3.1, 2.15],
        font_size=8.3,
    )
    add_manual_page_break(doc)

    heading(doc, "3. Windows PowerShell 快速开始", 1)
    heading(doc, "3.1 从哪个目录开始", 2)
    add_body(
        doc,
        "所有下面的相对路径都以包含 pyproject.toml、src、examples 和 .venv 的项目根目录为起点。当前机器上的完整目录是：",
    )
    add_code_block(doc, project_root_display)
    add_code_block(
        doc,
        r"""Set-Location -LiteralPath '<PROJECT_ROOT>'
Get-Location
Test-Path -LiteralPath '.\pyproject.toml'
Test-Path -LiteralPath '.\.venv\Scripts\simlab.exe'""".replace(
            "<PROJECT_ROOT>", project_root_display
        ),
    )
    add_note_box(
        doc,
        "成功判定",
        "Get-Location 应显示项目根目录，两个 Test-Path 都应返回 True。提示符形如 PS D:\\文件\\SimPy\\simpy-kpi-lab\\simpy-kpi-lab>。",
        kind="tip",
    )

    heading(doc, "3.2 首次安装", 2)
    add_code_block(
        doc,
        r"""py -3 -m venv .venv
.\.venv\Scripts\python.exe -m pip install -e \".[dev]\"

.\.venv\Scripts\simlab.exe validate .\examples\service_center.yaml
.\.venv\Scripts\python.exe -m pytest
.\.venv\Scripts\ruff.exe check .""",
    )
    add_body(
        doc,
        "如果系统没有 py 启动器，可把第一行替换为已安装 Python 3.11+ 的绝对路径。安装采用 editable 模式，因此修改 src/simlab 下的代码后通常不必重新安装包。",
    )

    heading(doc, "3.3 日常运行", 2)
    add_code_block(
        doc,
        r"""# 只校验，不执行仿真
.\.venv\Scripts\simlab.exe validate .\examples\service_center.yaml

# 串行执行：最容易排查问题
.\.venv\Scripts\simlab.exe run .\examples\service_center.yaml --workers 1

# 并行执行：workers 为进程数
.\.venv\Scripts\simlab.exe run .\examples\service_center.yaml --workers 4

# 临时覆盖输出目录
.\.venv\Scripts\simlab.exe run .\examples\service_center.yaml --output .\outputs\my_run""",
    )
    add_body(
        doc,
        "示例参数网格有 3 个场景，每场景 20 次 replication，因此终端会显示“开始执行 60 次仿真”。workers 只改变任务执行方式，不应改变固定 seed 下的 replication 数值。",
    )

    heading(doc, "3.4 CLI 命令速查", 2)
    add_table(
        doc,
        ["命令", "用途", "关键选项"],
        [
            ("simlab validate CONFIG", "校验 YAML，并报告场景数与总任务数", "无网络；不需要 API Key"),
            ("simlab run CONFIG", "运行场景 × replication 并保存结果", "--workers、--output"),
            ("simlab run CONFIG --analyze", "先保存本地结果，再请求 AI 分析", "--question、--model"),
            ("simlab analyze RESULTS", "只分析已有 results.json，不重跑仿真", "--question、--model、--output"),
        ],
        [2.25, 2.65, 1.6],
        font_size=8.5,
    )

    heading(doc, "4. YAML 配置详解", 1)
    add_body(
        doc,
        "配置根节点包含 project_name、simulation、experiment 和 openai。所有模型均设置 extra=\"forbid\"，字段拼写错误或未知字段会直接校验失败，避免静默采用错误假设。",
    )

    heading(doc, "4.1 simulation：系统模型", 2)
    add_table(
        doc,
        ["字段", "类型 / 默认值", "含义与约束"],
        [
            ("name", "str / service_system", "仿真模型名称"),
            ("until", "float / 必填", "仿真终止时刻；必须 > 0"),
            ("warmup", "float / 0", "删除期终点；0 ≤ warmup < until"),
            ("first_arrival_at_zero", "bool / true", "true 时首个实体在 t=0 到达；否则先采样一次间隔"),
            ("max_arrivals", "int|null / null", "可选到达上限；如设置必须 ≥ 1"),
            ("arrival_interarrival", "Distribution / 必填", "相邻到达的时间间隔；必须具有严格正支持"),
            ("stations", "list / 至少 1 个", "实体按列表顺序依次访问；工位名必须唯一"),
            ("cycle_time_target", "float|null / null", "服务水平阈值；设置时必须 > 0"),
        ],
        [1.72, 1.42, 3.36],
        font_size=8.15,
    )

    heading(doc, "4.2 station：工位", 2)
    add_table(
        doc,
        ["字段", "约束", "说明"],
        [
            ("name", "非空且全局唯一", "同时用于 KPI 路径和随机流名称；尽量使用稳定、简短的 ASCII 名称"),
            ("capacity", "整数，默认 1，且 ≥ 1", "simpy.Resource 的并行服务槽数量"),
            ("service_time", "Distribution", "每次开始服务时从该工位独立随机流抽样"),
        ],
        [1.35, 2.0, 3.15],
        font_size=8.5,
    )

    heading(doc, "4.3 Distribution：支持的分布", 2)
    add_table(
        doc,
        ["kind", "参数", "校验规则", "示例"],
        [
            ("exponential", "mean", "mean > 0", "{kind: exponential, mean: 5}"),
            ("deterministic", "value", "value ≥ 0；到达间隔不能为 0", "{kind: deterministic, value: 5}"),
            ("uniform", "low, high", "0 ≤ low ≤ high；到达分布 low 不能为 0", "{kind: uniform, low: 3, high: 7}"),
            ("triangular", "low, mode, high", "0 ≤ low ≤ mode ≤ high；到达分布 low 不能为 0", "{kind: triangular, low: 3, mode: 5, high: 9}"),
        ],
        [1.1, 1.28, 2.2, 1.92],
        font_size=7.9,
    )
    add_note_box(
        doc,
        "时间单位",
        "框架不强制“分钟”或“小时”，但一份配置中的到达、服务、warmup、until、周期目标和全部时间 KPI 必须使用同一单位。示例统一按分钟理解。",
        kind="warn",
    )

    heading(doc, "4.4 experiment：多次实验", 2)
    add_table(
        doc,
        ["字段", "默认值", "含义"],
        [
            ("replications", "10", "每个场景独立重复次数；必须 ≥ 1"),
            ("base_seed", "20260825", "所有派生随机流的根 seed；必须 ≥ 0"),
            ("common_random_numbers", "true", "不同场景同编号 replication 是否共享随机输入基线"),
            ("confidence_level", "0.95", "正态近似置信区间水平；必须在 0 与 1 之间"),
            ("parameter_grid", "{}", "相对 simulation 的点路径 → 候选值列表；列表不能为空"),
            ("output_dir", "outputs", "结果目录；--output 可临时覆盖"),
        ],
        [1.88, 1.1, 3.52],
        font_size=8.3,
    )

    heading(doc, "4.5 openai：可选 AI 分析", 2)
    add_table(
        doc,
        ["字段", "默认值 / 约束", "作用"],
        [
            ("model", "gpt-5.6", "Responses API 模型；可被 SIMLAB_OPENAI_MODEL 或 --model 覆盖"),
            ("max_output_tokens", "2500；≥ 256", "结构化分析的最大输出 token"),
            ("timeout_seconds", "60；> 0", "OpenAI 客户端请求超时秒数"),
            ("max_retries", "2；0–10", "客户端自动重试次数"),
            ("store", "false", "是否允许 API 存储响应；示例显式关闭"),
        ],
        [1.72, 1.55, 3.23],
        font_size=8.3,
    )
    add_manual_page_break(doc)

    heading(doc, "5. 示例配置逐段说明", 1)
    add_body(
        doc,
        "下面是项目随附的完整 examples/service_center.yaml。它产生 3 个 specialist 容量场景，每个场景 20 次，共 60 次仿真。",
    )
    add_code_block(
        doc,
        """project_name: service-center-capacity-study

simulation:
  name: two-stage-service-center
  until: 480
  warmup: 60
  first_arrival_at_zero: true
  arrival_interarrival:
    kind: exponential
    mean: 5.0
  stations:
    - name: registration
      capacity: 1
      service_time:
        kind: exponential
        mean: 3.0
    - name: specialist
      capacity: 2
      service_time:
        kind: triangular
        low: 4.0
        mode: 8.0
        high: 15.0
  cycle_time_target: 30

experiment:
  replications: 20
  base_seed: 20260825
  common_random_numbers: true
  confidence_level: 0.95
  output_dir: outputs/service_center
  parameter_grid:
    stations.1.capacity: [2, 3, 4]

openai:
  model: gpt-5.6
  max_output_tokens: 2500
  timeout_seconds: 60
  max_retries: 2
  store: false""",
        font_size=7.75,
    )

    heading(doc, "5.1 参数路径与场景笛卡尔积", 2)
    add_body(
        doc,
        "parameter_grid 的路径相对于 simulation。stations.1.capacity 表示 stations 列表中下标 1（第二个工位）的 capacity。也允许写 simulation.stations.1.capacity。多个路径会做笛卡尔积。",
    )
    add_code_block(
        doc,
        """parameter_grid:
  stations.0.capacity: [1, 2]
  arrival_interarrival.mean: [4.0, 5.0, 6.0]

# 场景数 = 2 × 3 = 6
# 总任务数 = 场景数 × replications""",
    )
    add_body(
        doc,
        "场景名称按展开顺序生成，例如 s001__stations_1_capacity=2。顺序取决于 YAML 中 parameter_grid 键和值的顺序；结果比较时应同时看 parameters 字段，而不要只依赖编号。",
    )
    add_note_box(
        doc,
        "常见校验错误",
        "路径不存在、列表下标越界、候选值类型不匹配、网格值列表为空，都会在 validate 或 run 加载阶段失败。先运行 validate 可以避免在长实验结束后才发现配置问题。",
        kind="warn",
    )

    heading(doc, "6. SimPy 事件机制", 1)
    heading(doc, "6.1 单次 replication 的生命周期", 2)
    add_numbered(
        doc,
        [
            "根据 seed 派生 arrivals 和 service:<station-name> 随机流。",
            "创建 simpy.Environment、每个工位的 simpy.Resource 和 KPICollector。",
            "启动 arrivals()；按 first_arrival_at_zero 决定首个到达时刻。",
            "每个 customer 依次进入所有工位队列，获得资源后抽样服务时间。",
            "每次排队、开工、服务区间和完成事件都通知 KPICollector。",
            "env.run(until=until) 在终止时刻停止，不执行恰好位于 until 的事件。",
            "finalize() 把事件记录转换为该次 replication 的 KPI 字典。",
        ],
    )

    heading(doc, "6.2 当前排队规则", 2)
    add_bullets(
        doc,
        [
            "工位是普通 simpy.Resource：先到先服务、无优先级、无抢占。",
            "实体不分类型；所有实体走相同串行路线。",
            "没有返工、放弃、批处理、资源故障、班次或动态容量。",
            "没有 drain 阶段；到达与在制实体都在 until 时刻被统一截断。",
            "max_arrivals 只限制产生的实体数，不会强制排空系统。",
        ],
    )
    add_note_box(
        doc,
        "建模含义",
        "这不是“通用流程图解释器”，而是一个清晰的串行服务网络模板。需要分流、优先级或故障时，应在 simulation.py 中扩展业务过程，同时补充相应事件、配置和测试。",
        kind="info",
    )

    heading(doc, "6.3 warmup 与统计窗口", 2)
    add_body(
        doc,
        "模型从空系统 t=0 启动，warmup 用于排除初始空系统对稳态指标的影响。统计窗口固定为 [warmup, until)。warmup 是人为设定的删除期，不是自动稳态检测；应通过试验检查 warmup 是否足够。",
    )
    add_table(
        doc,
        ["事件 / KPI", "纳入规则"],
        [
            ("arrivals", "到达时刻 warmup ≤ t < until"),
            ("completed", "完成时刻 warmup ≤ t < until；可能包含 warmup 前到达者"),
            ("周期 cohort", "到达时刻 ≥ warmup 且完成时刻 < until"),
            ("工位等待样本", "进入该工位队列时刻 ≥ warmup 且开工时刻 < until"),
            ("利用率 / 平均队长", "服务或排队区间与 [warmup, until) 的重叠时间"),
        ],
        [2.0, 4.5],
        font_size=8.6,
    )

    heading(doc, "7. 随机数、可复现性与共同随机数", 1)
    heading(doc, "7.1 命名随机流", 2)
    add_body(
        doc,
        "rng.py 使用带个性化字符串 simlab-rng-v1 的 BLAKE2b，把 base_seed 与命名空间稳定映射为 64 位整数。它不使用 Python hash()，因此不会受进程哈希随机化影响。",
    )
    add_code_block(
        doc,
        """base_seed
└─ replication:<编号>                  # replication seed
   ├─ arrivals                         # 到达流
   ├─ service:registration             # 登记服务流
   └─ service:specialist               # 专家服务流""",
    )
    add_body(
        doc,
        "工位子流按工位名称派生，而不是按列表位置派生。增加另一个工位不会直接改变已有工位的 seed；但系统状态和随机抽样次数可能变化，因此不能保证每个顾客仍消费相同的随机数。",
    )

    heading(doc, "7.2 common_random_numbers（CRN）", 2)
    add_table(
        doc,
        ["设置", "seed 命名空间", "适合场景", "注意"],
        [
            ("true", "只含 replication 编号", "容量、服务参数等相近方案比较", "同编号场景共享随机输入基线，可降低差值噪声"),
            ("false", "再加入规范化参数 JSON", "场景结构差异很大，或希望输入独立", "每个场景获得不同 seed"),
        ],
        [0.75, 1.78, 2.0, 1.97],
        font_size=8.1,
    )
    add_note_box(
        doc,
        "当前统计边界",
        "CRN 已用于生成共同随机输入，但 summary.csv 仍给出每个场景自己的边际均值和 CI；当前没有计算 candidate − baseline 的配对差值，也没有配对差值置信区间。",
        kind="warn",
    )

    heading(doc, "7.3 可复现检查清单", 2)
    add_bullets(
        doc,
        [
            "固定同一份 YAML、base_seed、依赖版本和 Python 版本。",
            "比较 replication 记录，而不要比较 generated_at 时间戳。",
            "workers=1 与 workers>1 应得到相同的 replication 和 summary 数值。",
            "修改工位名称会改变该工位的随机流；把名称视为模型契约的一部分。",
            "记录所用配置、代码版本与依赖锁定信息，才能进行长期审计。",
        ],
    )
    add_manual_page_break(doc)

    heading(doc, "8. KPI 口径与解读", 1)
    add_body(
        doc,
        "KPICollector 同时保存计数、样本列表和时间面积。metric_catalog 为每项可跨 replication 汇总的指标登记 role、direction、unit 与定义；capacity、service_starts、observed_cycle_count 等元数据不会进入 summary。",
    )

    heading(doc, "8.1 系统级 KPI", 2)
    system_kpis = [
        ("arrivals", "背景", "窗口内到达数", "需求规模；不直接判断好坏"),
        ("completed", "背景", "窗口内完成数", "可能包含 warmup 前到达者"),
        ("throughput_per_time_unit", "主要", "completed ÷ (until − warmup)", "越高通常越好"),
        ("wip_end", "护栏", "结束时仍在系统中的全部实体", "越低通常越好；注意包含 warmup 前到达者"),
        ("avg_cycle_time", "驱动", "已完成周期 cohort 的平均到达→完成时间", "越低通常越好"),
        ("p50_cycle_time", "驱动", "已完成周期 cohort 周期时间 P50", "典型体验"),
        ("p95_cycle_time", "主要", "已完成周期 cohort 周期时间 P95", "尾部体验，越低越好"),
        ("avg_wait_time", "驱动", "所有有效工位访问的访问加权平均等待", "不是每个实体平均总等待"),
        ("avg_total_wait_time", "驱动", "已完成周期 cohort 每实体跨工位总等待的平均", "衡量端到端排队负担"),
        ("service_level", "主要", "cycle ≤ cycle_time_target 的已完成 cohort 比例", "未设目标或无样本时为 null"),
        ("cycle_completion_fraction", "数据质量", "已完成周期 cohort 数 ÷ 窗口内到达数", "越接近 1，删失风险通常越低"),
        ("censored_cycle_count", "数据质量", "窗口内到达但结束前未完成数", "越高越需警惕完成案例偏差"),
    ]
    add_table(
        doc,
        ["指标", "角色", "计算口径", "使用提示"],
        system_kpis,
        [1.72, 0.72, 2.45, 1.61],
        font_size=7.55,
    )

    heading(doc, "8.2 工位级 KPI", 2)
    add_table(
        doc,
        ["指标路径", "角色", "计算口径", "解读"],
        [
            ("station.<name>.utilization", "驱动", "窗口内忙碌服务器时间面积 ÷ (capacity × 窗口长)", "过高常伴随排队，但不存在通用最优阈值"),
            ("station.<name>.avg_queue_length", "驱动", "队列人数时间面积 ÷ 窗口长", "含 until 时仍在排队实体的面积"),
            ("station.<name>.avg_wait_time", "驱动", "有效工位访问的平均等待", "越低通常越好"),
            ("station.<name>.p95_wait_time", "护栏", "有效工位访问等待时间 P95", "关注少数长等待"),
        ],
        [2.05, 0.72, 2.45, 1.28],
        font_size=7.8,
    )

    heading(doc, "8.3 百分位算法", 2)
    add_body(
        doc,
        "每次 replication 内先对样本排序，位置取 (n−1)×p；若位置落在两个样本之间，按小数部分做线性插值。随后汇总的是“各 replication 的 P95”这一标量，而不是把所有实体样本池化后再算一个 P95。",
    )

    heading(doc, "8.4 右删失与完成案例偏差", 2)
    add_note_box(
        doc,
        "必须一起看",
        "周期时间、总等待和服务水平只来自结束前完成的 cohort。靠近 until 到达且等待很久的实体更容易未完成，因此这些完成案例指标可能偏乐观。判断结果时必须同时检查 cycle_completion_fraction、censored_cycle_count 和 wip_end。",
        kind="danger",
    )
    add_body(
        doc,
        "当前没有停止到达后继续排空系统的 drain 阶段。若数据质量指标不好，可延长 until、降低 warmup 占比、减少到达强度，或在后续版本实现 stop-arrivals + drain。",
    )

    heading(doc, "9. 多次 replication 与统计汇总", 1)
    heading(doc, "9.1 为什么不能只跑一次", 2)
    add_body(
        doc,
        "单次随机仿真只是一个可能世界。多次 replication 使用不同 seed 独立重复同一场景，才能估计随机不确定性。框架把每个 replication 的 KPI 当作统计样本，不把同一次仿真中的每个顾客误当成独立实验。",
    )

    heading(doc, "9.2 summary.csv 的统计字段", 2)
    add_table(
        doc,
        ["字段", "定义", "备注"],
        [
            ("n", "该指标的有效 replication 数", "等同 n_valid"),
            ("n_total", "该场景计划执行的 replication 总数", "每个 catalog KPI 都输出"),
            ("n_missing", "n_total − n", "null 不会补 0"),
            ("mean", "有效 replication KPI 的算术平均", "主要点估计"),
            ("std", "样本标准差，分母 n−1", "n≤1 时为 null"),
            ("standard_error", "std ÷ √n", "n≤1 时为 null"),
            ("ci_low / ci_high", "mean ± z × standard_error", "正态近似；不截断比例的 0–1 边界"),
            ("min / max", "replication KPI 的最小/最大", "用于发现异常波动"),
        ],
        [1.6, 2.75, 2.15],
        font_size=8.2,
    )
    add_code_block(doc, "CI = mean ± NormalDist⁻¹((1 + confidence_level) / 2) × std / √n")
    add_note_box(
        doc,
        "置信区间不是“结果一定落在这里”",
        "当前 CI 描述在重复相同实验程序时均值估计的不确定性。replication 较少、分布重尾或 KPI 接近比例边界时，正态近似可能较粗糙；应增加 replication，并考虑 Student-t 或 bootstrap。",
        kind="warn",
    )

    heading(doc, "9.3 当前示例结果怎么读", 2)
    add_body(
        doc,
        "下表来自本地 outputs/service_center/summary.csv（20 次 replication/场景）。数值仅用于演示读法，不构成通用容量建议；示例没有成本指标，因此不能仅凭服务指标断言容量 4 一定是业务最优。",
    )
    add_table(
        doc,
        ["specialist 容量", "P95 周期时间 均值 [95% CI]", "服务水平 均值 [95% CI]", "吞吐率 均值 [95% CI]"],
        [
            ("2", "47.49 [41.64, 53.34]", "54.79% [44.77%, 64.80%]", "0.1986 [0.1914, 0.2057]"),
            ("3", "29.13 [26.68, 31.59]", "92.34% [88.64%, 96.04%]", "0.2036 [0.1959, 0.2112]"),
            ("4", "27.59 [25.24, 29.93]", "95.37% [92.74%, 98.01%]", "0.2030 [0.1950, 0.2109]"),
        ],
        [1.0, 1.95, 1.85, 1.7],
        font_size=7.65,
    )
    add_bullets(
        doc,
        [
            "容量 2 的 P95 周期时间明显高于 30 分钟目标，且服务水平较低。",
            "容量 2 增至 3 后，尾部周期与服务水平明显改善，吞吐仅小幅变化；主要收益是缓解拥堵。",
            "容量 3 与 4 的吞吐 CI 高度重叠；继续扩容前需结合成本、目标门槛、specialist 利用率、等待 P95、删失比例与 WIP。",
        ],
    )

    add_manual_page_break(doc)
    heading(doc, "10. 输出文件与结果结构", 1)
    add_body(doc, "run 完成后，输出目录至少包含三类本地结果；只有成功执行 AI 分析后才出现两个 ai_analysis 文件。")
    add_code_block(
        doc,
        """outputs/service_center/
├─ results.json          # 完整配置、随机流说明、指标目录、每次运行、汇总
├─ replications.csv      # 每个 replication 一行的宽表
├─ summary.csv           # 每个场景 × KPI 一行的统计汇总
├─ ai_analysis.json      # 可选：Pydantic 校验后的结构化分析
└─ ai_analysis.md        # 可选：面向阅读的 Markdown 分析""",
    )
    add_table(
        doc,
        ["文件", "适合用途", "重要注意"],
        [
            ("results.json", "完整审计、二次处理、独立 analyze", "schema_version 当前为 1.1；文件较大"),
            ("replications.csv", "统计复核、绘图、异常 replication 排查", "含 capacity、service_starts 等数值元数据；不是 summary"),
            ("summary.csv", "场景 KPI 比较、报表、AI 输入核心", "只包含 metric_catalog 登记 KPI"),
            ("ai_analysis.json", "程序读取 AI 结论", "固定文件名；再次分析会覆盖"),
            ("ai_analysis.md", "人工阅读或纳入报告", "AI 建议不是统计事实"),
        ],
        [1.55, 2.5, 2.45],
        font_size=8.2,
    )

    heading(doc, "10.1 results.json 顶层结构", 2)
    add_code_block(
        doc,
        """{
  \"schema_version\": \"1.1\",
  \"generated_at\": \"...UTC...\",
  \"project_name\": \"...\",
  \"config\": { ... },
  \"random_streams\": {
    \"method\": \"blake2b_namespaced_v1\",
    \"common_random_numbers\": true
  },
  \"metric_catalog\": [ ... ],
  \"replications\": [ ... ],
  \"summary\": [ ... ]
}""",
    )
    add_body(
        doc,
        "replications 中每条记录包含 scenario、parameters、replication、seed 和 metrics。summary 中每条记录包含场景、参数、指标元数据、有效/缺失样本数、均值、标准差、标准误、CI 与极值。",
    )

    heading(doc, "10.2 覆盖与归档", 2)
    add_note_box(
        doc,
        "输出目录会复用固定文件名",
        "再次运行到同一 output_dir 会覆盖 results.json、replications.csv 和 summary.csv；再次分析会覆盖 ai_analysis.json 与 ai_analysis.md。重要实验应使用新的 --output 目录，或在运行后立即归档整个目录。",
        kind="danger",
    )

    heading(doc, "11. OpenAI API 接入与 API Key", 1)
    add_body(
        doc,
        "项目接入的是 OpenAI Responses API，用它生成类似 ChatGPT 的 KPI 解读；官方产品接口名称不是一个单独的“ChatGPT API”。普通 validate 和 run 完全不需要 API Key，也不会联网。",
    )

    heading(doc, "11.1 获取 API Key", 2)
    add_numbered(
        doc,
        [
            "登录 OpenAI Platform（不是仅打开 ChatGPT 对话页面）。",
            "进入 API Keys 页面，选择合适的项目并创建新的 secret key。",
            "创建时复制并安全保存；密钥通常只会完整显示一次。",
            "确认该 API 项目有可用额度、计费设置和所选模型权限。",
        ],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("API Keys：")
    set_run_font(r, bold=True)
    add_hyperlink(p, "https://platform.openai.com/api-keys", "https://platform.openai.com/api-keys")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("官方 Quickstart：")
    set_run_font(r, bold=True)
    add_hyperlink(p, "https://developers.openai.com/api/docs/quickstart", "https://developers.openai.com/api/docs/quickstart")

    heading(doc, "11.2 在当前 PowerShell 会话中安全设置", 2)
    add_code_block(
        doc,
        r"""Set-Location -LiteralPath '<PROJECT_ROOT>'

# 输入时不回显，也不把真实 key 写进命令历史
$secureKey = Read-Host '粘贴 OpenAI API Key' -AsSecureString
$env:OPENAI_API_KEY = [System.Net.NetworkCredential]::new('', $secureKey).Password

# 只检查变量是否存在，不打印密钥
if (-not (Test-Path Env:OPENAI_API_KEY)) { throw 'OPENAI_API_KEY 未设置' }

# 直接分析已经生成的结果，不必重跑 60 次仿真
.\.venv\Scripts\simlab.exe analyze .\outputs\service_center\results.json

# 使用完后从当前会话删除
Remove-Item Env:OPENAI_API_KEY -ErrorAction SilentlyContinue""".replace(
            "<PROJECT_ROOT>", project_root_display
        ),
        font_size=7.85,
    )
    add_note_box(
        doc,
        "变量名必须完全一致",
        "应写 OPENAI_API_KEY，不要写反斜杠，不要把示例中的“...”当作真实 key。$env: 只对当前 PowerShell 进程及其子进程生效；新开终端后需要重新设置。",
        kind="warn",
    )
    add_bullets(
        doc,
        [
            "不要把 key 写入 YAML、Python 源码、截图、聊天消息、日志或提交到版本库。",
            "不要通过 Get-ChildItem Env: 或 Write-Host 输出完整 key。",
            "团队或生产环境应使用专用 secret manager 与最小权限项目 key。",
            "如怀疑泄露，应立即在 Platform 撤销并轮换 key。",
        ],
    )

    heading(doc, "11.3 运行分析", 2)
    add_code_block(
        doc,
        r"""# 仿真完成后立即分析
.\.venv\Scripts\simlab.exe run .\examples\service_center.yaml --analyze

# 带业务问题
.\.venv\Scripts\simlab.exe analyze .\outputs\service_center\results.json `
  --question '比较服务水平、P95 周期时间与工位利用率，并指出权衡。'

# 临时覆盖模型
.\.venv\Scripts\simlab.exe analyze .\outputs\service_center\results.json `
  --model gpt-5.6""",
    )
    add_body(
        doc,
        "模型优先级：--model 最高；run 加载 YAML 时，SIMLAB_OPENAI_MODEL 可覆盖 YAML；独立 analyze 默认使用 results.json 中已保存的 openai.model。官方模型目录在 2026-08-26 显示 gpt-5.6 是 GPT-5.6 Sol 的别名，但实际可用性仍取决于 API 项目权限。",
    )

    heading(doc, "11.4 API 实际接收什么", 2)
    add_table(
        doc,
        ["会发送", "不会发送"],
        [
            ("project_name", "OPENAI_API_KEY 本身"),
            ("simulation 与 experiment 配置", "replications 原始明细"),
            ("random_streams", "本地源代码与其他磁盘文件"),
            ("metric_catalog 与 summary", "未显式放入 payload 的数据"),
            ("用户的 --question", "AI 对本地 KPI 数值的修改权限"),
        ],
        [3.25, 3.25],
        font_size=8.5,
    )
    add_body(
        doc,
        "ai.py 使用 client.responses.parse(..., text_format=KPIAnalysis)，要求返回 executive_summary、best_scenario、findings、recommendations 和 caveats。Pydantic 校验成功后，才写入 JSON 和 Markdown。store 默认 false。",
    )
    p = doc.add_paragraph()
    r = p.add_run("Structured Outputs 官方指南：")
    set_run_font(r, bold=True)
    add_hyperlink(p, "https://developers.openai.com/api/docs/guides/structured-outputs", "https://developers.openai.com/api/docs/guides/structured-outputs")

    heading(doc, "11.5 API 错误排查", 2)
    add_table(
        doc,
        ["现象", "常见原因", "处理"],
        [
            ("OPENAI_API_KEY is not set", "变量未在当前进程中设置", "同一 PowerShell 会话重新设置；用 Test-Path 检查"),
            ("认证失败 / 401", "示例值、无效、已撤销或粘贴错误", "重新创建并设置；不要公开 key"),
            ("403 / 模型无权限", "项目权限或模型访问受限", "检查项目与模型权限；用 --model 选可访问模型"),
            ("429 / 额度不足", "速率、项目限额或余额", "稍后重试并检查 Limits / Billing"),
            ("请求超时", "网络慢或 60 秒不足", "检查网络；谨慎提高 timeout_seconds"),
            ("无法连接", "代理、防火墙、DNS 或 TLS", "检查网络环境；对已有 results.json 重试 analyze"),
            ("HTTP 5xx", "服务端临时错误", "客户端按 max_retries 重试；稍后重试"),
            ("未包含可解析分析", "拒答、输出不完整或结构化解析失败", "简化问题、检查 token 上限与模型支持"),
        ],
        [1.65, 2.2, 2.65],
        font_size=7.65,
    )
    add_note_box(
        doc,
        "针对你遇到的认证失败",
        "仿真结果已经成功写入 outputs/service_center，因此不必再次运行 60 次。获取并设置真实 API Key 后，直接执行 simlab analyze .\\outputs\\service_center\\results.json 即可。",
        kind="tip",
    )
    add_manual_page_break(doc)

    heading(doc, "12. 测试、验证与质量门禁", 1)
    add_body(
        doc,
        "当前项目已在本工作区验证：13 个 pytest 测试通过、ruff 通过、pip check 通过；同一配置的 60 次仿真在 workers=1 与 workers=2 下得到一致的 replication 与 summary 数值。generated_at 时间戳不同，因此完整 results.json 不应做字节级比较。",
    )
    add_code_block(
        doc,
        r"""# 单元与集成测试
.\.venv\Scripts\python.exe -m pytest

# 代码风格和静态规则
.\.venv\Scripts\ruff.exe check .

# 依赖一致性
.\.venv\Scripts\python.exe -m pip check""",
    )
    add_table(
        doc,
        ["测试文件", "覆盖重点"],
        [
            ("tests/test_config.py", "严格配置校验、分布与路径约束"),
            ("tests/test_rng.py", "稳定 seed 派生与命名空间"),
            ("tests/test_simulation.py", "单次仿真、KPI 和确定性行为"),
            ("tests/test_experiment.py", "场景展开、汇总、保存、串并行一致性"),
            ("tests/test_ai.py", "结构化分析、缺 key 与 API 错误映射"),
        ],
        [2.05, 4.45],
        font_size=8.6,
    )

    heading(doc, "12.1 每次正式实验前", 2)
    add_bullets(
        doc,
        [
            "先 validate，确认场景数和总任务数符合预期。",
            "用较小 replications 和独立 output 目录做 smoke run。",
            "检查时间单位、warmup、until、到达负荷与服务容量是否合理。",
            "明确主要 KPI、护栏、目标阈值与成本约束。",
            "重要模型修改后至少跑 pytest、ruff 和串并行一致性检查。",
        ],
    )

    heading(doc, "12.2 结果验收", 2)
    add_bullets(
        doc,
        [
            "确认每场景 n_total 等于配置的 replications。",
            "确认关键 KPI 的 n_missing 为 0 或解释缺失原因。",
            "同时查看 cycle_completion_fraction、censored_cycle_count 与 wip_end。",
            "查看 CI 宽度；如果结论依赖小差异，应增加 replication。",
            "检查利用率和排队指标是否符合流程直觉，并用简单手算场景做 sanity check。",
            "将原始配置、输出和代码版本一起归档。",
        ],
    )

    heading(doc, "13. 当前局限与建议增强", 1)
    add_table(
        doc,
        ["主题", "当前状态", "建议增强（尚未实现）"],
        [
            ("右删失", "until 直接截断，无 drain", "停止到达后排空；分离时间持续型窗口与 cohort 完成窗口"),
            ("置信区间", "正态近似边际 CI", "Student-t、bootstrap；比例边界处理"),
            ("CRN 比较", "只共享随机输入", "按 replication 计算 candidate−baseline 配对差值及 CI"),
            ("实验精度", "固定 replication 次数", "自适应 replication 与精度停止准则"),
            ("流程能力", "单类实体、串行 FIFO", "路由、优先级、抢占、返工、故障、班次与动态容量"),
            ("优化", "无成本 KPI 和自动搜索", "成本/收益指标、约束优化或试验设计层"),
            ("持久化", "固定文件名覆盖", "原子写入、manifest、运行 ID、断点续跑"),
            ("AI 可追溯性", "保存分析正文", "记录实际模型、response/request ID、usage 与输入哈希"),
        ],
        [1.18, 2.15, 3.17],
        font_size=7.85,
    )

    heading(doc, "13.1 扩展模型的推荐步骤", 2)
    add_numbered(
        doc,
        [
            "先写清业务事件、实体状态、资源规则和 KPI 口径，再改代码。",
            "在 config.py 增加严格字段与跨字段校验，保持 extra=\"forbid\"。",
            "在 simulation.py 中实现过程逻辑；为每个新随机过程增加独立命名随机流。",
            "在 KPICollector 中用事件或时间面积计算数值，并登记 metric_catalog。",
            "补充可手算的确定性 golden test、边界时刻测试和串并行一致性测试。",
            "先做小规模 smoke run，再增加 replication；AI 层最后接入。",
        ],
    )
    add_note_box(
        doc,
        "保持职责边界",
        "业务事件决定什么发生；KPICollector 决定如何测量；ExperimentRunner 决定如何重复和汇总；AI 只解释已经计算好的结果。这个边界是项目可验证、可审计的核心。",
        kind="info",
    )

    heading(doc, "14. 常见问题（FAQ）", 1)
    faq_rows = [
        ("PowerShell 应从哪里开始？", "从包含 pyproject.toml、src、examples、.venv 的项目根目录开始；当前是 D:\\文件\\SimPy\\simpy-kpi-lab\\simpy-kpi-lab。"),
        ("普通仿真需要 API Key 吗？", "不需要。只有 run --analyze 或 analyze 才调用 OpenAI API。"),
        ("认证失败后要重跑仿真吗？", "不用。results.json 已先保存，设置正确 key 后直接 analyze。"),
        ("为什么 60 次？", "3 个容量候选值 × 每场景 20 次 replication。"),
        ("workers 越大越好吗？", "不一定。任务小或 CPU/内存有限时并行开销可能抵消收益；先用 1 排错。"),
        ("为什么服务水平是 null？", "未设置 cycle_time_target，或该 replication 没有完成的周期 cohort。"),
        ("为什么某些 CI 是 null？", "有效 replication ≤ 1，无法计算样本标准差与标准误。"),
        ("利用率越高越好吗？", "没有通用答案。过高常增加等待；应与吞吐、队长、等待尾部和成本一起看。"),
        ("可以直接把 AI 建议当结论吗？", "不可以。AI 是解释层；本地 KPI、CI、业务约束与成本才是决策依据。"),
        ("如何避免覆盖旧结果？", "每次用 --output 指定带日期或实验名的新目录，并归档配置。"),
    ]
    add_table(doc, ["问题", "答案"], faq_rows, [2.25, 4.25], font_size=8.15)

    heading(doc, "15. 术语表", 1)
    add_table(
        doc,
        ["术语", "解释"],
        [
            ("离散事件仿真", "系统状态只在到达、开工、完成等离散事件发生时改变的仿真方法。"),
            ("replication", "同一场景、不同随机 seed 的一次完整独立运行。"),
            ("scenario", "参数网格中一组具体参数组合。"),
            ("warmup", "为降低空系统初始偏差而丢弃的早期时间段。"),
            ("cohort", "按到达/完成条件定义、用于周期指标的一组实体。"),
            ("右删失", "仿真结束时尚未观察到实体完成，因此其完整周期时间未知。"),
            ("CRN", "Common Random Numbers；让方案共享随机输入以提高差值比较精度。"),
            ("KPI catalog", "指标名称、角色、方向、单位与定义的机器可读目录。"),
            ("Structured Outputs", "让模型输出遵循给定 JSON Schema，并由本项目 Pydantic 模型解析。"),
        ],
        [1.8, 4.7],
        font_size=8.4,
    )

    heading(doc, "附录 A：关键路径与官方参考", 1)
    add_table(
        doc,
        ["类型", "路径 / 链接", "用途"],
        [
            ("项目根目录", project_root_display, "全部相对命令的起点"),
            ("示例配置", r"examples\service_center.yaml", "两阶段容量实验"),
            ("本地结果", r"outputs\service_center\results.json", "认证失败后直接 analyze 的输入"),
            ("项目说明", r"README.md", "快速入门与设计摘要"),
            ("OpenAI API Key", "https://platform.openai.com/api-keys", "创建与撤销密钥"),
            ("OpenAI Quickstart", "https://developers.openai.com/api/docs/quickstart", "环境变量与首个请求"),
            ("Structured Outputs", "https://developers.openai.com/api/docs/guides/structured-outputs", "结构化响应模式"),
            ("模型目录", "https://developers.openai.com/api/docs/models", "当前模型 ID 与能力"),
        ],
        [1.3, 3.65, 1.55],
        font_size=7.65,
    )

    heading(doc, "附录 B：最短可执行流程", 1)
    add_code_block(
        doc,
        r"""Set-Location -LiteralPath '<PROJECT_ROOT>'
# 1) 校验
.\.venv\Scripts\simlab.exe validate .\examples\service_center.yaml
# 2) 本地仿真（不需要 API Key）
.\.venv\Scripts\simlab.exe run .\examples\service_center.yaml --workers 1
# 3) 需要 AI 解读时，再临时设置真实 key
$secureKey = Read-Host '粘贴 OpenAI API Key' -AsSecureString
$env:OPENAI_API_KEY = [System.Net.NetworkCredential]::new('', $secureKey).Password
.\.venv\Scripts\simlab.exe analyze .\outputs\service_center\results.json
Remove-Item Env:OPENAI_API_KEY -ErrorAction SilentlyContinue""".replace(
            "<PROJECT_ROOT>", project_root_display
        ),
        font_size=7.2,
        trailing_space=False,
    )
    terminal = doc.add_paragraph()
    terminal.paragraph_format.space_before = Pt(0)
    terminal.paragraph_format.space_after = Pt(0)
    terminal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    terminal.paragraph_format.line_spacing = Pt(1)
    return doc


def main() -> None:
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
